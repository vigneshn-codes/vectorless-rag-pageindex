import io
from typing import List
import logging

from backend.parsers.base import BaseParser, ParsedPage

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """Parses DOCX files using python-docx. Extracts paragraphs and tables."""

    @property
    def supported_extensions(self) -> List[str]:
        return [".docx", ".doc"]

    def parse(self, file_bytes: bytes, filename: str) -> List[ParsedPage]:
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx is not installed.")
            return []

        pages: List[ParsedPage] = []
        try:
            buf = io.BytesIO(file_bytes)
            doc = Document(buf)

            text_blocks = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_blocks.append(text)

            for table in doc.tables:
                table_lines = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        table_lines.append(" | ".join(cells))
                if table_lines:
                    text_blocks.append("\n".join(table_lines))

            if not text_blocks:
                return []

            # Group text blocks into ~800-char chunks
            TARGET_SIZE = 800
            current_chunk: List[str] = []
            current_len = 0
            chunk_index = 0

            for block in text_blocks:
                block_len = len(block)
                if current_len + block_len > TARGET_SIZE and current_chunk:
                    content = "\n\n".join(current_chunk)
                    pages.append(
                        ParsedPage(
                            page_number=None,
                            content=content,
                            metadata={"chunk_index": chunk_index},
                        )
                    )
                    chunk_index += 1
                    current_chunk = [block]
                    current_len = block_len
                else:
                    current_chunk.append(block)
                    current_len += block_len

            if current_chunk:
                content = "\n\n".join(current_chunk)
                pages.append(
                    ParsedPage(
                        page_number=None,
                        content=content,
                        metadata={"chunk_index": chunk_index},
                    )
                )

        except Exception as e:
            logger.error(f"Failed to parse DOCX {filename}: {e}")

        return pages
